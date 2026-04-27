"""
Document parser: extracts and chunks text from PDF and DOCX resumes.
"""

import io
import logging
from typing import List

import PyPDF2
from docx import Document

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract all text from a PDF file given its raw bytes."""
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n".join(pages).strip()
    except Exception as exc:
        logger.error("PDF extraction failed: %s", exc)
        raise ValueError(f"Could not read PDF: {exc}") from exc


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract all paragraph text from a DOCX file given its raw bytes."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also pull text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs).strip()
    except Exception as exc:
        logger.error("DOCX extraction failed: %s", exc)
        raise ValueError(f"Could not read DOCX: {exc}") from exc


def extract_resume_text(uploaded_file) -> str:
    """
    High-level entry point: detect format from filename and dispatch
    to the appropriate extractor.
    """
    file_bytes = uploaded_file.read()
    name = uploaded_file.name.lower()

    if name.endswith(".pdf"):
        text = extract_text_from_pdf(file_bytes)
    elif name.endswith(".docx"):
        text = extract_text_from_docx(file_bytes)
    else:
        raise ValueError(
            "Unsupported file type. Please upload a PDF or DOCX file."
        )

    if not text:
        raise ValueError(
            "No text could be extracted from the file. "
            "Make sure the document is not scanned or image-only."
        )

    logger.info("Extracted %d characters from %s", len(text), uploaded_file.name)
    return text


def chunk_text(
    text: str,
    chunk_size: int = 400,
    overlap: int = 60,
) -> List[str]:
    """
    Split *text* into overlapping word-based chunks suitable for embedding.

    Args:
        text: Source text.
        chunk_size: Target number of words per chunk.
        overlap: Number of words shared between consecutive chunks.

    Returns:
        List of non-empty string chunks.
    """
    words = text.split()
    if not words:
        return []

    if len(words) <= chunk_size:
        return [text]

    chunks: List[str] = []
    step = chunk_size - overlap
    start = 0

    while start < len(words):
        end = start + chunk_size
        chunk = " ".join(words[start:end])
        if chunk.strip():
            chunks.append(chunk)
        start += step

    return chunks
